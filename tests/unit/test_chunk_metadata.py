"""Unit tests for chunk metadata filtering."""

from __future__ import annotations

from pathlib import Path

import docx as python_docx

from src.core.constants import CHUNK_SECTION_KEY, LAYOUT_DOCUMENT_METADATA_KEYS
from src.domain.entities.document import Document
from src.infrastructure.loaders.docx_loader import DocxLoader
from src.infrastructure.loaders.markdown_loader import MarkdownLoader
from src.rag.chunking.contextual_headers import ContextualHeadersChunker, prepend_headers
from src.rag.chunking.metadata import chunk_metadata
from src.rag.chunking.recursive_chunker import RecursiveChunker


class TestChunkMetadata:
    def test_layout_keys_constant(self) -> None:
        assert frozenset({"tables", "figures", "sections", "headings", "slides"}) == (
            LAYOUT_DOCUMENT_METADATA_KEYS
        )

    def test_chunk_metadata_excludes_layout_keys(self) -> None:
        metadata = {
            "loader": "docling",
            "filename": "report.pdf",
            CHUNK_SECTION_KEY: "Intro",
            "tables": [{"table_id": "table-1"}],
            "figures": [{"figure_id": "figure-1"}],
            "sections": ["Intro"],
            "headings": ["Intro"],
            "slides": [{"title": "Intro", "text": "body"}],
        }
        filtered = chunk_metadata(metadata)
        assert filtered == {
            "loader": "docling",
            "filename": "report.pdf",
            CHUNK_SECTION_KEY: "Intro",
        }

    def test_chunk_metadata_promotes_sections_to_section_key(self) -> None:
        filtered = chunk_metadata(
            {
                "loader": "docx",
                "sections": ["Introduction", "Details"],
            }
        )
        assert filtered == {
            "loader": "docx",
            CHUNK_SECTION_KEY: "Introduction",
        }

    def test_chunk_metadata_promotes_headings_to_section_key(self) -> None:
        filtered = chunk_metadata(
            {
                "loader": "markdown",
                "headings": ["Title", "Section One"],
            }
        )
        assert filtered == {
            "loader": "markdown",
            CHUNK_SECTION_KEY: "Title",
        }

    def test_chunk_metadata_prefers_existing_section_key(self) -> None:
        filtered = chunk_metadata(
            {
                CHUNK_SECTION_KEY: "Revenue",
                "sections": ["Intro"],
                "headings": ["Title"],
            }
        )
        assert filtered[CHUNK_SECTION_KEY] == "Revenue"

    def test_chunk_metadata_skips_empty_section_outlines(self) -> None:
        filtered = chunk_metadata({"loader": "docx", "sections": [], "headings": []})
        assert CHUNK_SECTION_KEY not in filtered

    def test_recursive_chunker_uses_filtered_metadata(self) -> None:
        document = Document(
            source="/tmp/report.pdf",
            content="Paragraph one.\n\nParagraph two.",
            metadata={
                "loader": "docling",
                "tables": [{"table_id": "table-1"}],
                "figures": [{"figure_id": "figure-1"}],
                "sections": ["Intro"],
                "headings": ["Intro"],
            },
        )
        chunks = RecursiveChunker(chunk_size=100, overlap=10).chunk(document)
        assert len(chunks) >= 1
        for chunk in chunks:
            assert "tables" not in chunk.metadata
            assert "figures" not in chunk.metadata
            assert "sections" not in chunk.metadata
            assert "headings" not in chunk.metadata
            assert chunk.metadata[CHUNK_SECTION_KEY] == "Intro"

    def test_docx_loader_path_preserves_section_in_contextual_headers(self, tmp_path: Path) -> None:
        path = tmp_path / "sample.docx"
        doc = python_docx.Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("This is the first paragraph.")
        doc.save(str(path))

        document = DocxLoader().load(path)
        chunker = ContextualHeadersChunker(RecursiveChunker(chunk_size=500))
        chunks = chunker.chunk(document)
        assert chunks
        header = prepend_headers(document, chunks[0])
        assert "Introduction" in header
        assert "—" not in header.split("Section:")[1].split("|")[0]

    def test_markdown_loader_path_preserves_section_in_contextual_headers(
        self, tmp_path: Path
    ) -> None:
        path = tmp_path / "sample.md"
        path.write_text("# Title\n\nFirst paragraph.\n", encoding="utf-8")

        document = MarkdownLoader().load(path)
        chunker = ContextualHeadersChunker(RecursiveChunker(chunk_size=500))
        chunks = chunker.chunk(document)
        assert chunks
        header = prepend_headers(document, chunks[0])
        assert "Title" in header
        assert "—" not in header.split("Section:")[1].split("|")[0]
