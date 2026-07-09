from __future__ import annotations

from pathlib import Path

import docx as python_docx

from src.core.constants import CHUNK_SECTION_KEY
from src.core.exceptions import DocumentLoadError
from src.domain.entities.document import Document


class DocxLoader:
    """Loads text from DOCX files using python-docx."""

    @staticmethod
    def load(path: Path) -> Document:
        try:
            doc = python_docx.Document(str(path))

            paragraphs = [p.text for p in doc.paragraphs]
            content = "\n\n".join(p for p in paragraphs if p.strip())

            section_titles: list[str] = []
            for p in doc.paragraphs:
                style = p.style
                if (
                    style is not None
                    and (style.name or "").startswith("Heading")
                    and p.text.strip()
                ):
                    section_titles.append(p.text)

            metadata: dict[str, object] = {
                "filename": path.name,
                "extension": path.suffix.lower(),
                "loader": "docx",
                "paragraph_count": len(paragraphs),
                "sections": section_titles,
            }
            if section_titles:
                metadata[CHUNK_SECTION_KEY] = section_titles[0]

            return Document(
                source=str(path.resolve()),
                content=content,
                metadata=metadata,
            )
        except DocumentLoadError:
            raise
        except Exception as exc:
            raise DocumentLoadError(f"Cannot load DOCX: {path}", cause=exc) from exc
